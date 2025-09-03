from typing import Dict, List, Any, Optional
import os
import json
import csv
import xml.etree.ElementTree as ET
from datetime import datetime
from jinja2 import Environment, FileSystemLoader, Template
import markdown
import logging

logger = logging.getLogger(__name__)

class DocumentationGenerator:
    """Generates formatted documentation from extracted database metadata."""
    
    def __init__(self, output_dir: str = "output"):
        self.output_dir = output_dir
        self.ensure_output_directory()
        
        # Initialize Jinja2 environment
        template_dir = os.path.join(os.path.dirname(__file__), "templates")
        if not os.path.exists(template_dir):
            os.makedirs(template_dir)
            self.create_default_templates()
        
        self.jinja_env = Environment(loader=FileSystemLoader(template_dir))
        self.jinja_env.filters['format_date'] = self._format_date
        self.jinja_env.filters['format_number'] = self._format_number
    
    def ensure_output_directory(self):
        """Ensure output directory exists."""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def generate_all_formats(self, documentation: Dict[str, Any], options: Dict[str, Any] = None) -> Dict[str, str]:
        """Generate documentation in all supported formats."""
        results = {}
        options = options or {}
        
        try:
            # Core formats
            if options.get('generate_html', True):
                results['html'] = self.generate_html_documentation(documentation)
            if options.get('generate_markdown', True):
                results['markdown'] = self.generate_markdown_documentation(documentation)
            if options.get('generate_json', True):
                results['json'] = self.generate_json_documentation(documentation)
            
            # Advanced formats
            advanced_results = self.generate_advanced_formats(documentation, options)
            results.update(advanced_results)
            
            logger.info("All documentation formats generated successfully")
            return results
        except Exception as e:
            logger.error(f"Failed to generate documentation: {str(e)}")
            raise
    
    def generate_html_documentation(self, documentation: Dict[str, Any]) -> str:
        """Generate comprehensive HTML documentation."""
        try:
            template = self.jinja_env.get_template('database_documentation.html')
            html_content = template.render(doc=documentation)
            
            output_file = os.path.join(self.output_dir, "database_documentation.html")
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            logger.info(f"HTML documentation generated: {output_file}")
            return output_file
        except Exception as e:
            logger.error(f"Failed to generate HTML documentation: {str(e)}")
            raise
    
    def generate_markdown_documentation(self, documentation: Dict[str, Any]) -> str:
        """Generate comprehensive Markdown documentation."""
        try:
            template = self.jinja_env.get_template('database_documentation.md')
            markdown_content = template.render(doc=documentation)
            
            output_file = os.path.join(self.output_dir, "database_documentation.md")
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            logger.info(f"Markdown documentation generated: {output_file}")
            return output_file
        except Exception as e:
            logger.error(f"Failed to generate Markdown documentation: {str(e)}")
            raise
    
    def generate_json_documentation(self, documentation: Dict[str, Any]) -> str:
        """Generate JSON documentation for programmatic access."""
        try:
            output_file = os.path.join(self.output_dir, "database_documentation.json")
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(documentation, f, indent=2, default=str)
            
            logger.info(f"JSON documentation generated: {output_file}")
            return output_file
        except Exception as e:
            logger.error(f"Failed to generate JSON documentation: {str(e)}")
            raise
    
    def create_default_templates(self):
        """Create default Jinja2 templates."""
        template_dir = os.path.join(os.path.dirname(__file__), "templates")
        
        # HTML Template
        html_template = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ doc.metadata.database_name }} - Database Documentation</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
        .header { background: #f4f4f4; padding: 20px; border-radius: 5px; margin-bottom: 30px; }
        .section { margin-bottom: 30px; }
        .table-container { overflow-x: auto; }
        table { border-collapse: collapse; width: 100%; margin-bottom: 20px; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        th { background-color: #f2f2f2; font-weight: bold; }
        .schema-section { margin-bottom: 40px; border: 1px solid #ddd; padding: 20px; border-radius: 5px; }
        .table-section { margin: 20px 0; padding: 15px; background: #f9f9f9; border-radius: 3px; }
        .column-list { margin: 10px 0; }
        .constraint-info { font-size: 0.9em; color: #666; margin: 5px 0; }
        .toc { background: #f9f9f9; padding: 20px; border-radius: 5px; margin-bottom: 30px; }
        .toc ul { list-style-type: none; padding-left: 20px; }
        .toc a { text-decoration: none; color: #007acc; }
        .toc a:hover { text-decoration: underline; }
        .stats-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin: 20px 0; }
        .stat-box { background: #f0f8ff; padding: 15px; border-radius: 5px; text-align: center; }
        .stat-number { font-size: 2em; font-weight: bold; color: #007acc; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ doc.metadata.database_name }} - Database Documentation</h1>
        <p><strong>Server:</strong> {{ doc.metadata.server_name }}</p>
        <p><strong>Generated:</strong> {{ doc.metadata.extraction_date | format_date }}</p>
        <p><strong>Size:</strong> {{ doc.metadata.used_mb | format_number }} MB used of {{ doc.metadata.size_mb | format_number }} MB allocated</p>
    </div>

    <div class="toc">
        <h2>Table of Contents</h2>
        <ul>
            <li><a href="#overview">Database Overview</a></li>
            <li><a href="#schemas">Schemas</a></li>
            <li><a href="#tables">Tables</a></li>
            <li><a href="#views">Views</a></li>
            <li><a href="#procedures">Stored Procedures</a></li>
            <li><a href="#functions">Functions</a></li>
            <li><a href="#relationships">Relationships</a></li>
        </ul>
    </div>

    <div id="overview" class="section">
        <h2>Database Overview</h2>
        <div class="stats-grid">
            <div class="stat-box">
                <div class="stat-number">{{ doc.statistics.total_schemas }}</div>
                <div>Schemas</div>
            </div>
            <div class="stat-box">
                <div class="stat-number">{{ doc.statistics.total_tables }}</div>
                <div>Tables</div>
            </div>
            <div class="stat-box">
                <div class="stat-number">{{ doc.statistics.total_views }}</div>
                <div>Views</div>
            </div>
            <div class="stat-box">
                <div class="stat-number">{{ doc.statistics.total_procedures }}</div>
                <div>Stored Procedures</div>
            </div>
            <div class="stat-box">
                <div class="stat-number">{{ doc.statistics.total_functions }}</div>
                <div>Functions</div>
            </div>
            <div class="stat-box">
                <div class="stat-number">{{ doc.statistics.total_rows | format_number }}</div>
                <div>Total Rows</div>
            </div>
        </div>
    </div>

    <div id="schemas" class="section">
        <h2>Schemas</h2>
        <div class="table-container">
            <table>
                <thead>
                    <tr>
                        <th>Schema Name</th>
                        <th>Principal</th>
                    </tr>
                </thead>
                <tbody>
                    {% for schema in doc.schemas %}
                    <tr>
                        <td>{{ schema.name }}</td>
                        <td>{{ schema.principal }}</td>
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
        </div>
    </div>

    <div id="tables" class="section">
        <h2>Tables</h2>
        {% for table in doc.tables %}
        <div class="table-section">
            <h3>{{ table.schema_name }}.{{ table.table_name }}</h3>
            {% if table.description %}
            <p><strong>Description:</strong> {{ table.description }}</p>
            {% endif %}
            <p><strong>Rows:</strong> {{ table.row_count | format_number }}</p>
            
            <h4>Columns</h4>
            <div class="table-container">
                <table>
                    <thead>
                        <tr>
                            <th>Column</th>
                            <th>Data Type</th>
                            <th>Nullable</th>
                            <th>Identity</th>
                            <th>Default</th>
                            <th>Description</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for column in table.columns %}
                        <tr>
                            <td>{{ column.name }}</td>
                            <td>{{ column.data_type }}</td>
                            <td>{{ 'Yes' if column.is_nullable else 'No' }}</td>
                            <td>{{ 'Yes' if column.is_identity else 'No' }}</td>
                            <td>{{ column.default_value }}</td>
                            <td>{{ column.description }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>

            {% if table.primary_keys %}
            <div class="constraint-info">
                <strong>Primary Key:</strong> {{ table.primary_keys | map(attribute='column_name') | join(', ') }}
            </div>
            {% endif %}

            {% if table.foreign_keys %}
            <div class="constraint-info">
                <strong>Foreign Keys:</strong>
                {% for fk in table.foreign_keys %}
                <div>{{ fk.parent_column }} → {{ fk.referenced_schema }}.{{ fk.referenced_table }}.{{ fk.referenced_column }}</div>
                {% endfor %}
            </div>
            {% endif %}

            {% if table.indexes %}
            <div class="constraint-info">
                <strong>Indexes:</strong>
                {% for idx in table.indexes %}
                <div>{{ idx.index_name }} ({{ idx.index_type }}): {{ idx.columns }}</div>
                {% endfor %}
            </div>
            {% endif %}
        </div>
        {% endfor %}
    </div>

    <div id="views" class="section">
        <h2>Views</h2>
        {% for view in doc.views %}
        <div class="table-section">
            <h3>{{ view.schema_name }}.{{ view.view_name }}</h3>
            {% if view.description %}
            <p><strong>Description:</strong> {{ view.description }}</p>
            {% endif %}
            
            <h4>Columns</h4>
            <div class="table-container">
                <table>
                    <thead>
                        <tr>
                            <th>Column</th>
                            <th>Data Type</th>
                            <th>Nullable</th>
                            <th>Description</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for column in view.columns %}
                        <tr>
                            <td>{{ column.name }}</td>
                            <td>{{ column.data_type }}</td>
                            <td>{{ 'Yes' if column.is_nullable else 'No' }}</td>
                            <td>{{ column.description }}</td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>
        </div>
        {% endfor %}
    </div>

    <div id="procedures" class="section">
        <h2>Stored Procedures</h2>
        <div class="table-container">
            <table>
                <thead>
                    <tr>
                        <th>Schema</th>
                        <th>Procedure Name</th>
                        <th>Created</th>
                        <th>Modified</th>
                        <th>Description</th>
                    </tr>
                </thead>
                <tbody>
                    {% for proc in doc.stored_procedures %}
                    <tr>
                        <td>{{ proc.schema_name }}</td>
                        <td>{{ proc.procedure_name }}</td>
                        <td>{{ proc.created | format_date }}</td>
                        <td>{{ proc.modified | format_date }}</td>
                        <td>{{ proc.description }}</td>
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
        </div>
    </div>

    <div id="functions" class="section">
        <h2>Functions</h2>
        <div class="table-container">
            <table>
                <thead>
                    <tr>
                        <th>Schema</th>
                        <th>Function Name</th>
                        <th>Type</th>
                        <th>Created</th>
                        <th>Modified</th>
                        <th>Description</th>
                    </tr>
                </thead>
                <tbody>
                    {% for func in doc.functions %}
                    <tr>
                        <td>{{ func.schema_name }}</td>
                        <td>{{ func.function_name }}</td>
                        <td>{{ func.type }}</td>
                        <td>{{ func.created | format_date }}</td>
                        <td>{{ func.modified | format_date }}</td>
                        <td>{{ func.description }}</td>
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
        </div>
    </div>

    <div id="relationships" class="section">
        <h2>Database Relationships</h2>
        <p>Total Foreign Key Relationships: {{ doc.relationships.relationship_count }}</p>
        
        <div class="table-container">
            <table>
                <thead>
                    <tr>
                        <th>Foreign Key</th>
                        <th>Parent Table</th>
                        <th>Parent Column</th>
                        <th>Referenced Table</th>
                        <th>Referenced Column</th>
                        <th>On Delete</th>
                        <th>On Update</th>
                    </tr>
                </thead>
                <tbody>
                    {% for fk in doc.relationships.foreign_keys %}
                    <tr>
                        <td>{{ fk.foreign_key_name }}</td>
                        <td>{{ fk.parent_schema }}.{{ fk.parent_table }}</td>
                        <td>{{ fk.parent_column }}</td>
                        <td>{{ fk.referenced_schema }}.{{ fk.referenced_table }}</td>
                        <td>{{ fk.referenced_column }}</td>
                        <td>{{ fk.on_delete }}</td>
                        <td>{{ fk.on_update }}</td>
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
        </div>
    </div>

    <footer style="margin-top: 50px; padding-top: 20px; border-top: 1px solid #ddd; font-size: 0.9em; color: #666;">
        <p>Generated on {{ doc.metadata.extraction_date | format_date }} using Azure SQL Database Documentation Tool</p>
    </footer>
</body>
</html>'''
        
        with open(os.path.join(template_dir, 'database_documentation.html'), 'w', encoding='utf-8') as f:
            f.write(html_template)
        
        # Markdown Template
        markdown_template = '''# {{ doc.metadata.database_name }} - Database Documentation

**Server:** {{ doc.metadata.server_name }}  
**Generated:** {{ doc.metadata.extraction_date | format_date }}  
**Size:** {{ doc.metadata.used_mb | format_number }} MB used of {{ doc.metadata.size_mb | format_number }} MB allocated

## Table of Contents
- [Database Overview](#database-overview)
- [Schemas](#schemas)
- [Tables](#tables)
- [Views](#views)
- [Stored Procedures](#stored-procedures)
- [Functions](#functions)
- [Relationships](#relationships)

## Database Overview

| Metric | Count |
|--------|-------|
| Schemas | {{ doc.statistics.total_schemas }} |
| Tables | {{ doc.statistics.total_tables }} |
| Views | {{ doc.statistics.total_views }} |
| Stored Procedures | {{ doc.statistics.total_procedures }} |
| Functions | {{ doc.statistics.total_functions }} |
| Total Rows | {{ doc.statistics.total_rows | format_number }} |

## Schemas

| Schema Name | Principal |
|-------------|-----------|
{% for schema in doc.schemas -%}
| {{ schema.name }} | {{ schema.principal }} |
{% endfor %}

## Tables

{% for table in doc.tables %}
### {{ table.schema_name }}.{{ table.table_name }}

{% if table.description %}
**Description:** {{ table.description }}
{% endif %}

**Rows:** {{ table.row_count | format_number }}

#### Columns

| Column | Data Type | Nullable | Identity | Default | Description |
|--------|-----------|----------|----------|---------|-------------|
{% for column in table.columns -%}
| {{ column.name }} | {{ column.data_type }} | {{ 'Yes' if column.is_nullable else 'No' }} | {{ 'Yes' if column.is_identity else 'No' }} | {{ column.default_value }} | {{ column.description }} |
{% endfor %}

{% if table.primary_keys %}
**Primary Key:** {{ table.primary_keys | map(attribute='column_name') | join(', ') }}
{% endif %}

{% if table.foreign_keys %}
**Foreign Keys:**
{% for fk in table.foreign_keys -%}
- {{ fk.parent_column }} → {{ fk.referenced_schema }}.{{ fk.referenced_table }}.{{ fk.referenced_column }}
{% endfor %}
{% endif %}

{% if table.indexes %}
**Indexes:**
{% for idx in table.indexes -%}
- {{ idx.index_name }} ({{ idx.index_type }}): {{ idx.columns }}
{% endfor %}
{% endif %}

{% endfor %}

## Views

{% for view in doc.views %}
### {{ view.schema_name }}.{{ view.view_name }}

{% if view.description %}
**Description:** {{ view.description }}
{% endif %}

#### Columns

| Column | Data Type | Nullable | Description |
|--------|-----------|----------|-------------|
{% for column in view.columns -%}
| {{ column.name }} | {{ column.data_type }} | {{ 'Yes' if column.is_nullable else 'No' }} | {{ column.description }} |
{% endfor %}

{% endfor %}

## Stored Procedures

| Schema | Procedure Name | Created | Modified | Description |
|--------|----------------|---------|----------|-------------|
{% for proc in doc.stored_procedures -%}
| {{ proc.schema_name }} | {{ proc.procedure_name }} | {{ proc.created | format_date }} | {{ proc.modified | format_date }} | {{ proc.description }} |
{% endfor %}

## Functions

| Schema | Function Name | Type | Created | Modified | Description |
|--------|---------------|------|---------|----------|-------------|
{% for func in doc.functions -%}
| {{ func.schema_name }} | {{ func.function_name }} | {{ func.type }} | {{ func.created | format_date }} | {{ func.modified | format_date }} | {{ func.description }} |
{% endfor %}

## Relationships

**Total Foreign Key Relationships:** {{ doc.relationships.relationship_count }}

| Foreign Key | Parent Table | Parent Column | Referenced Table | Referenced Column | On Delete | On Update |
|-------------|--------------|---------------|------------------|-------------------|-----------|-----------|
{% for fk in doc.relationships.foreign_keys -%}
| {{ fk.foreign_key_name }} | {{ fk.parent_schema }}.{{ fk.parent_table }} | {{ fk.parent_column }} | {{ fk.referenced_schema }}.{{ fk.referenced_table }} | {{ fk.referenced_column }} | {{ fk.on_delete }} | {{ fk.on_update }} |
{% endfor %}

---
*Generated on {{ doc.metadata.extraction_date | format_date }} using Azure SQL Database Documentation Tool*'''
        
        with open(os.path.join(template_dir, 'database_documentation.md'), 'w', encoding='utf-8') as f:
            f.write(markdown_template)
        
        logger.info("Default templates created")
    
    def _format_date(self, date_str: str) -> str:
        """Format ISO date string for display."""
        if not date_str:
            return ""
        try:
            dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            return dt.strftime('%Y-%m-%d %H:%M:%S')
        except:
            return date_str
    
    def _format_number(self, number: Any) -> str:
        """Format number with thousands separators."""
        if number is None:
            return "0"
        try:
            return f"{int(float(number)):,}"
        except:
            return str(number)
    
    def generate_advanced_formats(self, documentation: Dict[str, Any], options: Dict[str, Any] = None) -> Dict[str, str]:
        """Generate documentation in advanced formats (PDF, Excel, CSV, XML, Word)."""
        options = options or {}
        results = {}
        
        try:
            if options.get('generate_pdf', False):
                results['pdf'] = self.generate_pdf_documentation(documentation, options)
            
            if options.get('generate_excel', False):
                results['excel'] = self.generate_excel_documentation(documentation, options)
            
            if options.get('generate_csv', False):
                results['csv'] = self.generate_csv_documentation(documentation, options)
            
            if options.get('generate_xml', False):
                results['xml'] = self.generate_xml_documentation(documentation, options)
            
            if options.get('generate_word', False):
                results['word'] = self.generate_word_documentation(documentation, options)
            
            logger.info("Advanced format documentation generated successfully")
            return results
            
        except Exception as e:
            logger.error(f"Failed to generate advanced format documentation: {str(e)}")
            raise
    
    def generate_pdf_documentation(self, documentation: Dict[str, Any], options: Dict[str, Any] = None) -> str:
        """Generate PDF documentation using HTML to PDF conversion."""
        try:
            # First generate HTML
            html_content = self.generate_html_documentation(documentation)
            
            # Try to use weasyprint or reportlab for PDF generation
            try:
                import weasyprint
                pdf_path = os.path.join(self.output_dir, "database_documentation.pdf")
                
                # Create PDF with custom styling for print
                pdf_css = self._get_pdf_css_styles(options)
                weasyprint.HTML(string=html_content).write_pdf(
                    pdf_path,
                    stylesheets=[weasyprint.CSS(string=pdf_css)]
                )
                
                logger.info(f"PDF documentation generated: {pdf_path}")
                return pdf_path
                
            except ImportError:
                # Fallback: create a text-based PDF using reportlab
                from reportlab.lib.pagesizes import letter, A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.lib import colors
                
                pdf_path = os.path.join(self.output_dir, "database_documentation.pdf")
                doc = SimpleDocTemplate(pdf_path, pagesize=A4)
                styles = getSampleStyleSheet()
                story = []
                
                # Title
                title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], 
                                           fontSize=24, spaceAfter=30)
                story.append(Paragraph(f"Database Documentation: {documentation['metadata']['database_name']}", title_style))
                story.append(Spacer(1, 20))
                
                # Overview
                story.append(Paragraph("Database Overview", styles['Heading2']))
                overview_data = [
                    ['Database Name', documentation['metadata']['database_name']],
                    ['Server', documentation['metadata']['server_name']],
                    ['Generated', documentation['metadata']['extraction_date']],
                    ['Size (MB)', f"{documentation['metadata']['size_mb']:.1f}"],
                    ['Total Tables', str(documentation['statistics']['total_tables'])],
                    ['Total Views', str(documentation['statistics']['total_views'])],
                    ['Total Procedures', str(documentation['statistics']['total_procedures'])]
                ]
                
                overview_table = Table(overview_data, colWidths=[2*inch, 3*inch])
                overview_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(overview_table)
                story.append(Spacer(1, 20))
                
                # Tables section
                if documentation.get('tables'):
                    story.append(Paragraph("Database Tables", styles['Heading2']))
                    
                    for table in documentation['tables'][:10]:  # Limit for PDF size
                        story.append(Paragraph(f"{table['schema_name']}.{table['table_name']}", styles['Heading3']))
                        
                        if table.get('description'):
                            story.append(Paragraph(f"Description: {table['description']}", styles['Normal']))
                        
                        story.append(Paragraph(f"Rows: {table.get('row_count', 0):,}", styles['Normal']))
                        story.append(Spacer(1, 10))
                
                doc.build(story)
                logger.info(f"PDF documentation generated using ReportLab: {pdf_path}")
                return pdf_path
                
        except Exception as e:
            logger.error(f"Failed to generate PDF documentation: {str(e)}")
            # Return HTML as fallback
            return self.generate_html_documentation(documentation)
    
    def generate_excel_documentation(self, documentation: Dict[str, Any], options: Dict[str, Any] = None) -> str:
        """Generate Excel workbook with multiple sheets for database documentation."""
        try:
            import xlsxwriter
            
            excel_path = os.path.join(self.output_dir, "database_documentation.xlsx")
            workbook = xlsxwriter.Workbook(excel_path)
            
            # Define formats
            header_format = workbook.add_format({
                'bold': True,
                'font_size': 12,
                'bg_color': '#366092',
                'font_color': 'white',
                'border': 1
            })
            
            cell_format = workbook.add_format({
                'border': 1,
                'text_wrap': True
            })
            
            number_format = workbook.add_format({
                'border': 1,
                'num_format': '#,##0'
            })
            
            # Overview Sheet
            overview_sheet = workbook.add_worksheet('Overview')
            overview_sheet.write('A1', 'Database Overview', header_format)
            
            overview_data = [
                ('Database Name', documentation['metadata']['database_name']),
                ('Server Name', documentation['metadata']['server_name']),
                ('Generated Date', documentation['metadata']['extraction_date']),
                ('Size (MB)', documentation['metadata']['size_mb']),
                ('Used Space (MB)', documentation['metadata']['used_mb']),
                ('Total Schemas', documentation['statistics']['total_schemas']),
                ('Total Tables', documentation['statistics']['total_tables']),
                ('Total Views', documentation['statistics']['total_views']),
                ('Total Procedures', documentation['statistics']['total_procedures']),
                ('Total Functions', documentation['statistics']['total_functions']),
                ('Total Rows', documentation['statistics']['total_rows'])
            ]
            
            for row, (label, value) in enumerate(overview_data, 2):
                overview_sheet.write(row, 0, label, header_format)
                if isinstance(value, (int, float)):
                    overview_sheet.write(row, 1, value, number_format)
                else:
                    overview_sheet.write(row, 1, str(value), cell_format)
            
            # Tables Sheet
            if documentation.get('tables'):
                tables_sheet = workbook.add_worksheet('Tables')
                table_headers = ['Schema', 'Table Name', 'Type', 'Created', 'Modified', 'Row Count', 'Description']
                
                for col, header in enumerate(table_headers):
                    tables_sheet.write(0, col, header, header_format)
                
                for row, table in enumerate(documentation['tables'], 1):
                    tables_sheet.write(row, 0, table['schema_name'], cell_format)
                    tables_sheet.write(row, 1, table['table_name'], cell_format)
                    tables_sheet.write(row, 2, table.get('type', ''), cell_format)
                    tables_sheet.write(row, 3, table.get('created', ''), cell_format)
                    tables_sheet.write(row, 4, table.get('modified', ''), cell_format)
                    tables_sheet.write(row, 5, table.get('row_count', 0), number_format)
                    tables_sheet.write(row, 6, table.get('description', ''), cell_format)
                
                tables_sheet.autofit()
            
            # Columns Sheet
            if documentation.get('tables'):
                columns_sheet = workbook.add_worksheet('Columns')
                column_headers = ['Schema', 'Table', 'Column', 'Data Type', 'Nullable', 'Default', 'Description']
                
                for col, header in enumerate(column_headers):
                    columns_sheet.write(0, col, header, header_format)
                
                row = 1
                for table in documentation['tables']:
                    for column in table.get('columns', []):
                        columns_sheet.write(row, 0, table['schema_name'], cell_format)
                        columns_sheet.write(row, 1, table['table_name'], cell_format)
                        columns_sheet.write(row, 2, column['name'], cell_format)
                        columns_sheet.write(row, 3, column['data_type'], cell_format)
                        columns_sheet.write(row, 4, 'Yes' if column['is_nullable'] else 'No', cell_format)
                        columns_sheet.write(row, 5, column.get('default_value', ''), cell_format)
                        columns_sheet.write(row, 6, column.get('description', ''), cell_format)
                        row += 1
                
                columns_sheet.autofit()
            
            # Views Sheet
            if documentation.get('views'):
                views_sheet = workbook.add_worksheet('Views')
                view_headers = ['Schema', 'View Name', 'Created', 'Modified', 'Description']
                
                for col, header in enumerate(view_headers):
                    views_sheet.write(0, col, header, header_format)
                
                for row, view in enumerate(documentation['views'], 1):
                    views_sheet.write(row, 0, view['schema_name'], cell_format)
                    views_sheet.write(row, 1, view['view_name'], cell_format)
                    views_sheet.write(row, 2, view.get('created', ''), cell_format)
                    views_sheet.write(row, 3, view.get('modified', ''), cell_format)
                    views_sheet.write(row, 4, view.get('description', ''), cell_format)
                
                views_sheet.autofit()
            
            # Procedures Sheet
            if documentation.get('stored_procedures'):
                procs_sheet = workbook.add_worksheet('Stored Procedures')
                proc_headers = ['Schema', 'Procedure Name', 'Created', 'Modified', 'Description']
                
                for col, header in enumerate(proc_headers):
                    procs_sheet.write(0, col, header, header_format)
                
                for row, proc in enumerate(documentation['stored_procedures'], 1):
                    procs_sheet.write(row, 0, proc['schema_name'], cell_format)
                    procs_sheet.write(row, 1, proc['procedure_name'], cell_format)
                    procs_sheet.write(row, 2, proc.get('created', ''), cell_format)
                    procs_sheet.write(row, 3, proc.get('modified', ''), cell_format)
                    procs_sheet.write(row, 4, proc.get('description', ''), cell_format)
                
                procs_sheet.autofit()
            
            workbook.close()
            logger.info(f"Excel documentation generated: {excel_path}")
            return excel_path
            
        except ImportError:
            logger.warning("xlsxwriter not available, falling back to CSV")
            return self.generate_csv_documentation(documentation, options)
        except Exception as e:
            logger.error(f"Failed to generate Excel documentation: {str(e)}")
            raise
    
    def generate_csv_documentation(self, documentation: Dict[str, Any], options: Dict[str, Any] = None) -> str:
        """Generate CSV files for database documentation."""
        try:
            csv_dir = os.path.join(self.output_dir, "csv_export")
            if not os.path.exists(csv_dir):
                os.makedirs(csv_dir)
            
            files_created = []
            
            # Tables CSV
            if documentation.get('tables'):
                tables_csv = os.path.join(csv_dir, "tables.csv")
                with open(tables_csv, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Schema', 'Table Name', 'Type', 'Created', 'Modified', 'Row Count', 'Description'])
                    
                    for table in documentation['tables']:
                        writer.writerow([
                            table['schema_name'],
                            table['table_name'],
                            table.get('type', ''),
                            table.get('created', ''),
                            table.get('modified', ''),
                            table.get('row_count', 0),
                            table.get('description', '')
                        ])
                files_created.append(tables_csv)
            
            # Columns CSV
            if documentation.get('tables'):
                columns_csv = os.path.join(csv_dir, "columns.csv")
                with open(columns_csv, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Schema', 'Table', 'Column', 'Data Type', 'Nullable', 'Identity', 'Default', 'Description'])
                    
                    for table in documentation['tables']:
                        for column in table.get('columns', []):
                            writer.writerow([
                                table['schema_name'],
                                table['table_name'],
                                column['name'],
                                column['data_type'],
                                column['is_nullable'],
                                column.get('is_identity', False),
                                column.get('default_value', ''),
                                column.get('description', '')
                            ])
                files_created.append(columns_csv)
            
            # Views CSV
            if documentation.get('views'):
                views_csv = os.path.join(csv_dir, "views.csv")
                with open(views_csv, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Schema', 'View Name', 'Created', 'Modified', 'Description'])
                    
                    for view in documentation['views']:
                        writer.writerow([
                            view['schema_name'],
                            view['view_name'],
                            view.get('created', ''),
                            view.get('modified', ''),
                            view.get('description', '')
                        ])
                files_created.append(views_csv)
            
            # Relationships CSV
            if documentation.get('relationships', {}).get('foreign_keys'):
                relationships_csv = os.path.join(csv_dir, "relationships.csv")
                with open(relationships_csv, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    writer.writerow(['Foreign Key', 'Parent Schema', 'Parent Table', 'Parent Column', 
                                   'Referenced Schema', 'Referenced Table', 'Referenced Column', 'On Delete', 'On Update'])
                    
                    for fk in documentation['relationships']['foreign_keys']:
                        writer.writerow([
                            fk['foreign_key_name'],
                            fk['parent_schema'],
                            fk['parent_table'],
                            fk['parent_column'],
                            fk['referenced_schema'],
                            fk['referenced_table'],
                            fk['referenced_column'],
                            fk['on_delete'],
                            fk['on_update']
                        ])
                files_created.append(relationships_csv)
            
            logger.info(f"CSV documentation generated in: {csv_dir}")
            return csv_dir
            
        except Exception as e:
            logger.error(f"Failed to generate CSV documentation: {str(e)}")
            raise
    
    def generate_xml_documentation(self, documentation: Dict[str, Any], options: Dict[str, Any] = None) -> str:
        """Generate XML documentation."""
        try:
            # Create root element
            root = ET.Element("DatabaseDocumentation")
            
            # Metadata
            metadata = ET.SubElement(root, "Metadata")
            for key, value in documentation['metadata'].items():
                elem = ET.SubElement(metadata, key.title().replace('_', ''))
                elem.text = str(value) if value is not None else ""
            
            # Statistics
            statistics = ET.SubElement(root, "Statistics")
            for key, value in documentation['statistics'].items():
                if key != 'largest_tables':  # Skip complex nested data for XML
                    elem = ET.SubElement(statistics, key.title().replace('_', ''))
                    elem.text = str(value) if value is not None else "0"
            
            # Tables
            if documentation.get('tables'):
                tables_elem = ET.SubElement(root, "Tables")
                for table in documentation['tables']:
                    table_elem = ET.SubElement(tables_elem, "Table")
                    table_elem.set("schema", table['schema_name'])
                    table_elem.set("name", table['table_name'])
                    
                    for key, value in table.items():
                        if key not in ['schema_name', 'table_name', 'columns', 'primary_keys', 'foreign_keys', 'indexes']:
                            if value is not None:
                                elem = ET.SubElement(table_elem, key.title().replace('_', ''))
                                elem.text = str(value)
                    
                    # Columns
                    if table.get('columns'):
                        columns_elem = ET.SubElement(table_elem, "Columns")
                        for column in table['columns']:
                            col_elem = ET.SubElement(columns_elem, "Column")
                            col_elem.set("name", column['name'])
                            col_elem.set("dataType", column['data_type'])
                            col_elem.set("nullable", str(column['is_nullable']))
                            if column.get('description'):
                                col_elem.set("description", column['description'])
            
            # Views
            if documentation.get('views'):
                views_elem = ET.SubElement(root, "Views")
                for view in documentation['views']:
                    view_elem = ET.SubElement(views_elem, "View")
                    view_elem.set("schema", view['schema_name'])
                    view_elem.set("name", view['view_name'])
                    if view.get('description'):
                        view_elem.set("description", view['description'])
            
            # Write XML to file
            xml_path = os.path.join(self.output_dir, "database_documentation.xml")
            tree = ET.ElementTree(root)
            ET.indent(tree, space="  ", level=0)
            tree.write(xml_path, encoding='utf-8', xml_declaration=True)
            
            logger.info(f"XML documentation generated: {xml_path}")
            return xml_path
            
        except Exception as e:
            logger.error(f"Failed to generate XML documentation: {str(e)}")
            raise
    
    def generate_word_documentation(self, documentation: Dict[str, Any], options: Dict[str, Any] = None) -> str:
        """Generate Word document documentation."""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(f"Database Documentation: {documentation['metadata']['database_name']}", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata section
            doc.add_heading('Database Overview', level=1)
            
            metadata_table = doc.add_table(rows=1, cols=2)
            metadata_table.style = 'Table Grid'
            hdr_cells = metadata_table.rows[0].cells
            hdr_cells[0].text = 'Property'
            hdr_cells[1].text = 'Value'
            
            metadata_items = [
                ('Database Name', documentation['metadata']['database_name']),
                ('Server Name', documentation['metadata']['server_name']),
                ('Generated Date', documentation['metadata']['extraction_date']),
                ('Size (MB)', f"{documentation['metadata']['size_mb']:.1f}"),
                ('Used Space (MB)', f"{documentation['metadata']['used_mb']:.1f}")
            ]
            
            for prop, value in metadata_items:
                row_cells = metadata_table.add_row().cells
                row_cells[0].text = prop
                row_cells[1].text = str(value)
            
            # Statistics
            doc.add_heading('Database Statistics', level=1)
            
            stats_para = doc.add_paragraph()
            stats_para.add_run(f"Total Schemas: ").bold = True
            stats_para.add_run(f"{documentation['statistics']['total_schemas']}\n")
            stats_para.add_run(f"Total Tables: ").bold = True
            stats_para.add_run(f"{documentation['statistics']['total_tables']}\n")
            stats_para.add_run(f"Total Views: ").bold = True
            stats_para.add_run(f"{documentation['statistics']['total_views']}\n")
            stats_para.add_run(f"Total Procedures: ").bold = True
            stats_para.add_run(f"{documentation['statistics']['total_procedures']}\n")
            stats_para.add_run(f"Total Functions: ").bold = True
            stats_para.add_run(f"{documentation['statistics']['total_functions']}\n")
            stats_para.add_run(f"Total Rows: ").bold = True
            stats_para.add_run(f"{documentation['statistics']['total_rows']:,}")
            
            # Tables section (limited for document size)
            if documentation.get('tables'):
                doc.add_heading('Database Tables (Sample)', level=1)
                
                for table in documentation['tables'][:5]:  # Limit to first 5 tables
                    doc.add_heading(f"{table['schema_name']}.{table['table_name']}", level=2)
                    
                    table_para = doc.add_paragraph()
                    if table.get('description'):
                        table_para.add_run(f"Description: ").bold = True
                        table_para.add_run(f"{table['description']}\n")
                    
                    table_para.add_run(f"Row Count: ").bold = True
                    table_para.add_run(f"{table.get('row_count', 0):,}\n")
                    
                    if table.get('columns'):
                        table_para.add_run(f"Columns: ").bold = True
                        table_para.add_run(f"{len(table['columns'])}")
            
            # Save document
            word_path = os.path.join(self.output_dir, "database_documentation.docx")
            doc.save(word_path)
            
            logger.info(f"Word documentation generated: {word_path}")
            return word_path
            
        except ImportError:
            logger.warning("python-docx not available, falling back to HTML")
            return self.generate_html_documentation(documentation)
        except Exception as e:
            logger.error(f"Failed to generate Word documentation: {str(e)}")
            raise
    
    def _get_pdf_css_styles(self, options: Dict[str, Any] = None) -> str:
        """Get CSS styles optimized for PDF generation."""
        return """
        @page {
            size: A4;
            margin: 1in;
        }
        
        body {
            font-family: 'Arial', sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #333;
        }
        
        h1 {
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            page-break-before: always;
        }
        
        h1:first-child {
            page-break-before: avoid;
        }
        
        h2 {
            color: #34495e;
            margin-top: 25px;
            border-bottom: 1px solid #bdc3c7;
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
            font-size: 10pt;
        }
        
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        
        th {
            background-color: #f8f9fa;
            font-weight: bold;
        }
        
        .stats-grid {
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 15px;
            margin: 20px 0;
        }
        
        .stat-box {
            border: 1px solid #ddd;
            padding: 15px;
            text-align: center;
            background-color: #f8f9fa;
        }
        
        .toc {
            page-break-after: always;
        }
        """
    
    def generate_batch_documentation(self, projects: List[Dict[str, Any]], options: Dict[str, Any] = None) -> Dict[str, List[str]]:
        """Generate documentation for multiple projects in batch."""
        results = {
            'successful': [],
            'failed': [],
            'files_created': []
        }
        
        options = options or {}
        batch_dir = options.get('batch_output_dir', os.path.join(self.output_dir, 'batch_export'))
        
        if not os.path.exists(batch_dir):
            os.makedirs(batch_dir)
        
        for project in projects:
            try:
                project_id = project.get('id', 'unknown')
                project_name = project.get('name', f'project_{project_id}')
                
                # Create project-specific output directory
                project_dir = os.path.join(batch_dir, f"{project_name}_{project_id}")
                if not os.path.exists(project_dir):
                    os.makedirs(project_dir)
                
                # Update output directory for this project
                original_output_dir = self.output_dir
                self.output_dir = project_dir
                
                # Generate documentation for this project
                project_docs = project.get('documentation', {})
                if project_docs:
                    if options.get('generate_html', True):
                        html_file = self.generate_html_documentation(project_docs)
                        results['files_created'].append(html_file)
                    
                    if options.get('generate_pdf', False):
                        pdf_file = self.generate_pdf_documentation(project_docs, options)
                        results['files_created'].append(pdf_file)
                    
                    if options.get('generate_excel', False):
                        excel_file = self.generate_excel_documentation(project_docs, options)
                        results['files_created'].append(excel_file)
                
                results['successful'].append(project_name)
                
                # Restore original output directory
                self.output_dir = original_output_dir
                
            except Exception as e:
                logger.error(f"Failed to generate documentation for project {project.get('name', project_id)}: {str(e)}")
                results['failed'].append(f"{project.get('name', project_id)}: {str(e)}")
        
        logger.info(f"Batch documentation completed. Successful: {len(results['successful'])}, Failed: {len(results['failed'])}")
        return results